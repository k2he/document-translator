"""
rebuild.py — Read translated_segments.json, patch translated text back into a
             copy of the DOCX, apply North American formatting, and write the
             output file.

Postprocessing: Normalizes punctuation spacing in final text after patching.
  - Removes spaces before punctuation
  - Ensures proper spacing after punctuation

Usage:
    python src/rebuild.py input/charpter-4.docx
"""

import json
import re
import shutil
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING

sys.path.insert(0, str(Path(__file__).resolve().parent))
from config import (
    INPUT_DIR,
    OUTPUT_DIR,
    WORK_DIR,
    TRANSLATED_SEGMENTS,
    PAGE_WIDTH_EMU,
    PAGE_HEIGHT_EMU,
    MARGIN_EMU,
    BODY_FONT,
    BODY_FONT_SIZE_PT,
    HEADING_CHAPTER_FONT,
    HEADING_CHAPTER_FONT_SIZE_PT,
    HEADING_CHAPTER_BOLD,
    HEADING_SECTION_FONT,
    HEADING_SECTION_FONT_SIZE_PT,
    HEADING_SECTION_BOLD,
    PARA_SPACING_AFTER_PT,
)


def normalize_punctuation_spacing(text: str) -> str:
    """
    Normalize English punctuation spacing in final translated text.
    
    Rules:
      1. Remove space(s) before punctuation (Western: ,.)
      2. Ensure exactly one space after punctuation if followed by a word
    
    Example:
      "chapters ,we clarified" → "chapters, we clarified"
      "slope ,  introducing" → "slope, introducing"
    """
    # Rule 1: Remove space(s) immediately before Western punctuation
    text = re.sub(r'\s+([,.])', r'\1', text)
    
    # Rule 2: Collapse multiple spaces after punctuation to single space
    text = re.sub(r'([,.])\s+', r'\1 ', text)
    
    # Rule 3: Ensure space after punctuation if followed by a non-punctuation word character
    text = re.sub(r'([,.])([a-zA-Z0-9])', r'\1 \2', text)
    
    return text


def _load_translations(path: Path) -> dict[str, str]:
    """Return {segment_id: translated_text} mapping."""
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    return {item["id"]: item["translation"] for item in data}


def _apply_page_format(doc: docx.Document) -> None:
    """Set US Letter page size and 1-inch margins on every section."""
    for section in doc.sections:
        section.page_width = PAGE_WIDTH_EMU
        section.page_height = PAGE_HEIGHT_EMU
        section.left_margin = MARGIN_EMU
        section.right_margin = MARGIN_EMU
        section.top_margin = MARGIN_EMU
        section.bottom_margin = MARGIN_EMU


def _apply_styles(doc: docx.Document) -> None:
    """
    Update built-in styles for body text, Heading 1 (chapter), and Heading 2
    (section) to match North American textbook conventions.

    Only touches styles that already exist in the document to avoid creating
    orphaned style entries.
    """
    style_updates = {
        "Normal": {
            "font_name": BODY_FONT,
            "font_size": Pt(BODY_FONT_SIZE_PT),
            "bold": None,  # None = don't touch
        },
        "Heading 1": {
            "font_name": HEADING_CHAPTER_FONT,
            "font_size": Pt(HEADING_CHAPTER_FONT_SIZE_PT),
            "bold": HEADING_CHAPTER_BOLD,
        },
        "Heading 2": {
            "font_name": HEADING_SECTION_FONT,
            "font_size": Pt(HEADING_SECTION_FONT_SIZE_PT),
            "bold": HEADING_SECTION_BOLD,
        },
    }

    existing_style_names = {s.name for s in doc.styles}

    for style_name, props in style_updates.items():
        if style_name not in existing_style_names:
            continue
        style = doc.styles[style_name]
        font = style.font
        font.name = props["font_name"]
        font.size = props["font_size"]
        if props["bold"] is not None:
            font.bold = props["bold"]

        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = Pt(BODY_FONT_SIZE_PT * 1.5)
        pf.space_after = Pt(PARA_SPACING_AFTER_PT)


def _normalize_all_paragraphs(doc: docx.Document) -> None:
    """
    Apply punctuation spacing normalization at the paragraph level.
    Combines all runs in a paragraph, normalizes, and reconstructs.
    This fixes spacing issues across run boundaries.
    """
    # Normalize body paragraphs
    for para in doc.paragraphs:
        _normalize_paragraph(para)
    
    # Normalize table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _normalize_paragraph(para)


# XML namespace URIs used for content-type guards
_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _has_manual_superscript(para) -> bool:
    """Return True if any run in the paragraph uses manual baseline-raise + small font
    to simulate a superscript exponent (e.g. x³, x¹⁰⁰, aˣ in this source doc).

    Detection heuristic: w:position >= 4 (raised ≥ 2 pt) AND w:sz <= 16 (font ≤ 8 pt).
    Normal body text never combines a large raise with a very small font size.
    """
    for r in para._element.findall(f'.//{{{_W_NS}}}r'):
        rPr = r.find(f'{{{_W_NS}}}rPr')
        if rPr is None:
            continue
        pos_el = rPr.find(f'{{{_W_NS}}}position')
        sz_el  = rPr.find(f'{{{_W_NS}}}sz')
        if pos_el is None or sz_el is None:
            continue
        try:
            pos = int(pos_el.get(f'{{{_W_NS}}}val', '0'))
            sz  = int(sz_el.get(f'{{{_W_NS}}}val', '99'))
        except ValueError:
            continue
        if pos >= 4 and sz <= 16:
            return True
    return False


def _has_protected_content(para) -> bool:
    """Return True if the paragraph contains content that must not be modified.

    Protected element types:
    - w:drawing       — DrawingML inline/anchored images (Format Picture)
    - w:pict          — VML picture container (Format Object, e.g. formula images)
    - m:oMath         — OMML math equations (Insert → Equation)
    - m:oMathPara     — OMML display math block
    - manual raised   — w:position >= 4 + w:sz <= 16 (hand-typeset exponents
                        like x³, x¹⁰⁰, aˣ using baseline shift + small font)
    """
    p = para._element
    return (
        p.find(f'.//{{{_W_NS}}}drawing') is not None
        or p.find(f'.//{{{_W_NS}}}pict') is not None
        or p.find(f'.//{{{_M_NS}}}oMath') is not None
        or p.find(f'.//{{{_M_NS}}}oMathPara') is not None
        or _has_manual_superscript(para)
    )


def _normalize_paragraph(para) -> None:
    """
    Normalize a single paragraph by combining all runs, normalizing text,
    and redistributing back to runs while preserving the run structure.
    Skips paragraphs that contain pictures or math formulas.
    """
    if not para.runs:
        return
    if _has_protected_content(para):
        return
    
    # Collect all text and normalize
    full_text = ''.join(run.text for run in para.runs)
    normalized_text = normalize_punctuation_spacing(full_text)
    
    # If no change, skip
    if normalized_text == full_text:
        return
    
    # Redistribute normalized text back to runs, preserving structure
    # Strategy: clear all runs except the first, then update text
    if para.runs:
        # Keep first run's formatting, update its text to full normalized
        first_run = para.runs[0]
        first_run.text = normalized_text
        
        # Remove all other runs by clearing their text
        for run in para.runs[1:]:
            # Remove run by clearing its element
            run._element.getparent().remove(run._element)


def _patch_paragraphs(paragraphs, translations: dict[str, list[dict]], location_prefix: str, offset: int = 0):
    """
    Patch runs in *paragraphs* using the translations index.

    *translations* is keyed by (location, para_index, run_index):
        { (location, para_index, run_index): translated_text }
    """
    for para_idx, para in enumerate(paragraphs):
        abs_idx = offset + para_idx
        for run_idx, run in enumerate(para.runs):
            key = (location_prefix, abs_idx, run_idx)
            if key in translations:
                run.text = translations[key]


def rebuild(docx_path: Path) -> Path:
    translations_raw = _load_translations(TRANSLATED_SEGMENTS)

    # Build a lookup keyed by (location, para_index, run_index)
    # We need the full segment info, so reload extracted_segments for position data.
    from config import EXTRACTED_SEGMENTS

    with open(EXTRACTED_SEGMENTS, encoding="utf-8") as f:
        segments = json.load(f)

    # Map segment_id → position
    position_map: dict[str, tuple[str, int, int]] = {
        seg["id"]: (seg["location"], seg["para_index"], seg["run_index"])
        for seg in segments
    }

    # Build (location, para_index, run_index) → translated_text
    patch_map: dict[tuple, str] = {}
    for seg_id, translation in translations_raw.items():
        if seg_id in position_map:
            loc, pi, ri = position_map[seg_id]
            patch_map[(loc, pi, ri)] = translation

    # Make a working copy so the original is never modified
    stem = docx_path.stem
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"{stem}_translated.docx"
    shutil.copy2(docx_path, output_path)

    doc = docx.Document(str(output_path))

    # Patch body paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            key = ("body", para_idx, run_idx)
            if key in patch_map:
                run.text = patch_map[key]

    # Patch table cells
    table_para_offset = len(doc.paragraphs)
    for tbl_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para_idx, para in enumerate(cell.paragraphs):
                    abs_idx = table_para_offset + para_idx
                    for run_idx, run in enumerate(para.runs):
                        key = (f"table_{tbl_idx}", abs_idx, run_idx)
                        if key in patch_map:
                            run.text = patch_map[key]
                table_para_offset += len(cell.paragraphs)

    # Apply punctuation spacing normalization to all paragraphs (body + tables)
    _normalize_all_paragraphs(doc)

    _apply_page_format(doc)
    _apply_styles(doc)

    doc.save(str(output_path))
    print(f"Rebuilt document → {output_path}", flush=True)
    return output_path


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Patch translations back into DOCX.")
    parser.add_argument("docx_file", help="Path to the original source DOCX file")
    args = parser.parse_args()

    docx_path = Path(args.docx_file)
    if not docx_path.exists():
        print(f"ERROR: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    if not TRANSLATED_SEGMENTS.exists():
        print(f"ERROR: Translated segments not found: {TRANSLATED_SEGMENTS}", file=sys.stderr)
        sys.exit(1)

    rebuild(docx_path)


if __name__ == "__main__":
    main()
