"""audit.py — AI audit helper for translated DOCX files.

Two sub-commands:

  extract  Walk the translated DOCX, collect all non-trivial English
           paragraphs, write work/audit_segments.json for AI review.

  apply    Read work/audited_segments.json (written by the Copilot skill)
           and patch the AI-fixed text back into the DOCX in-place.

Usage:
    python src/audit.py extract output/charpter-4_translated.docx
    python src/audit.py apply   output/charpter-4_translated.docx
"""

import argparse
import json
import sys
from pathlib import Path

import docx

sys.path.insert(0, str(Path(__file__).resolve().parent))
from config import WORK_DIR, AUDIT_SEGMENTS, AUDITED_SEGMENTS, TRANSLATION_STYLE

# Only audit paragraphs that have enough English text to be meaningful
_MIN_ALPHA_CHARS = 3
_MIN_TEXT_LEN = 8


def _is_auditable(text: str) -> bool:
    """Return True if the paragraph is English prose worth auditing."""
    stripped = text.strip()
    if len(stripped) < _MIN_TEXT_LEN:
        return False
    alpha_chars = sum(1 for c in stripped if c.isalpha())
    return alpha_chars >= _MIN_ALPHA_CHARS


def extract_for_audit(docx_path: Path) -> None:
    """
    Extract all auditable paragraphs from the translated DOCX and write
    work/audit_segments.json for the AI to review.
    """
    doc = docx.Document(str(docx_path))
    segments = []
    seg_id = 0

    # Body paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        text = "".join(run.text for run in para.runs)
        if _is_auditable(text):
            segments.append({
                "id": f"audit_{seg_id:05d}",
                "location": "body",
                "para_index": para_idx,
                "text": text,
            })
            seg_id += 1

    # Table cell paragraphs
    table_para_offset = len(doc.paragraphs)
    for tbl_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para_idx, para in enumerate(cell.paragraphs):
                    text = "".join(run.text for run in para.runs)
                    if _is_auditable(text):
                        segments.append({
                            "id": f"audit_{seg_id:05d}",
                            "location": f"table_{tbl_idx}",
                            "para_index": table_para_offset + para_idx,
                            "text": text,
                        })
                        seg_id += 1
                table_para_offset += len(cell.paragraphs)

    WORK_DIR.mkdir(parents=True, exist_ok=True)
    payload = {
        "style_instruction": TRANSLATION_STYLE,
        "segments": segments,
    }
    with open(AUDIT_SEGMENTS, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    style_note = f" | style: '{TRANSLATION_STYLE[:60]}…'" if TRANSLATION_STYLE else ""
    print(f"Extracted {len(segments)} paragraph(s) for AI audit → {AUDIT_SEGMENTS}{style_note}", flush=True)


def apply_audit(docx_path: Path) -> None:
    """
    Read work/audited_segments.json written by the Copilot skill and patch
    the AI-fixed text back into the DOCX in-place.

    Only paragraphs where 'fixed' differs from 'original' are updated.
    """
    if not AUDIT_SEGMENTS.exists():
        print(f"ERROR: {AUDIT_SEGMENTS} not found. Run extract first.", file=sys.stderr)
        sys.exit(1)
    if not AUDITED_SEGMENTS.exists():
        print(f"ERROR: {AUDITED_SEGMENTS} not found. Run AI audit phase first.", file=sys.stderr)
        sys.exit(1)

    with open(AUDIT_SEGMENTS, encoding="utf-8") as f:
        _audit_data = json.load(f)
    # Support both old list format and new {style_instruction, segments} format
    if isinstance(_audit_data, dict):
        audit_segs = _audit_data["segments"]
    else:
        audit_segs = _audit_data

    with open(AUDITED_SEGMENTS, encoding="utf-8") as f:
        audited_segs = json.load(f)

    # Build {id → fixed_text} for paragraphs that were changed by the AI
    fixes: dict[str, str] = {}
    for item in audited_segs:
        fixed = item.get("fixed", "").strip()
        original = item.get("original", "").strip()
        if fixed and fixed != original:
            fixes[item["id"]] = item["fixed"]

    if not fixes:
        print("AI audit made no changes — document unchanged.", flush=True)
        return

    # Build {(location, para_index) → fixed_text}
    patch_map: dict[tuple, str] = {}
    for seg in audit_segs:
        if seg["id"] in fixes:
            patch_map[(seg["location"], seg["para_index"])] = fixes[seg["id"]]

    doc = docx.Document(str(docx_path))

    # Patch body paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        key = ("body", para_idx)
        if key in patch_map:
            _replace_paragraph_text(para, patch_map[key])

    # Patch table cell paragraphs
    table_para_offset = len(doc.paragraphs)
    for tbl_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para_idx, para in enumerate(cell.paragraphs):
                    key = (f"table_{tbl_idx}", table_para_offset + para_idx)
                    if key in patch_map:
                        _replace_paragraph_text(para, patch_map[key])
                table_para_offset += len(cell.paragraphs)

    doc.save(str(docx_path))
    print(f"Applied {len(fixes)} AI audit fix(es) → {docx_path}", flush=True)


# XML namespace URIs used for content-type guards
_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _has_manual_superscript(para) -> bool:
    """Return True if any run in the paragraph uses manual baseline-raise + small font
    to simulate a superscript exponent (e.g. x³, x¹⁰⁰, aˣ in this source doc).

    Detection heuristic: w:position >= 4 (raised ≥ 2 pt) AND w:sz <= 16 (font ≤ 8 pt).
    Normal body text never combines a large raise with a very small font size.
    """
    for r in para._element.findall(f'.//{{{_W_NS}}}r'):
        rPr = r.find(f'{{{_W_NS}}}rPr')
        if rPr is None:
            continue
        pos_el = rPr.find(f'{{{_W_NS}}}position')
        sz_el  = rPr.find(f'{{{_W_NS}}}sz')
        if pos_el is None or sz_el is None:
            continue
        try:
            pos = int(pos_el.get(f'{{{_W_NS}}}val', '0'))
            sz  = int(sz_el.get(f'{{{_W_NS}}}val', '99'))
        except ValueError:
            continue
        if pos >= 4 and sz <= 16:
            return True
    return False


def _has_protected_content(para) -> bool:
    """Return True if the paragraph contains content that must not be modified.

    Protected element types:
    - w:drawing       — DrawingML inline/anchored images (Format Picture)
    - w:pict          — VML picture container (Format Object, e.g. formula images)
    - m:oMath         — OMML math equations (Insert → Equation)
    - m:oMathPara     — OMML display math block
    - manual raised   — w:position >= 4 + w:sz <= 16 (hand-typeset exponents
                        like x³, x¹⁰⁰, aˣ using baseline shift + small font)
    """
    p = para._element
    return (
        p.find(f'.//{{{_W_NS}}}drawing') is not None
        or p.find(f'.//{{{_W_NS}}}pict') is not None
        or p.find(f'.//{{{_M_NS}}}oMath') is not None
        or p.find(f'.//{{{_M_NS}}}oMathPara') is not None
        or _has_manual_superscript(para)
    )


def _replace_paragraph_text(para, new_text: str) -> None:
    """
    Replace all run text in a paragraph with new_text.
    Puts the full text in the first run and removes the rest,
    preserving the formatting of the first run.
    Skips paragraphs that contain pictures or math formulas.
    """
    if not para.runs:
        return
    if _has_protected_content(para):
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)


def main():
    parser = argparse.ArgumentParser(
        description="AI audit helper: extract paragraphs for AI review and apply fixes back."
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    extract_parser = subparsers.add_parser(
        "extract", help="Extract English paragraphs from DOCX for AI audit"
    )
    extract_parser.add_argument("docx_file", help="Path to the translated DOCX file")

    apply_parser = subparsers.add_parser(
        "apply", help="Patch AI-audited fixes back into the translated DOCX"
    )
    apply_parser.add_argument("docx_file", help="Path to the translated DOCX file")

    args = parser.parse_args()
    docx_path = Path(args.docx_file)

    if not docx_path.exists():
        print(f"ERROR: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    if args.command == "extract":
        extract_for_audit(docx_path)
    elif args.command == "apply":
        apply_audit(docx_path)


if __name__ == "__main__":
    main()
