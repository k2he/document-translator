"""
extract.py — Walk a DOCX file, detect Chinese text runs, and write
             work/extracted_segments.json for translation.

Usage:
    python src/extract.py input/charpter-4.docx
    python src/extract.py input/charpter-4.docx --test --pages 5
"""

import argparse
import json
import re
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

# Add repo root to path so config is importable from anywhere
sys.path.insert(0, str(Path(__file__).resolve().parent))
from config import INPUT_DIR, WORK_DIR, EXTRACTED_SEGMENTS

# Unicode range covering CJK Unified Ideographs (basic block)
_CJK_RE = re.compile(r"[\u4e00-\u9fff\u3400-\u4dbf\uff00-\uffef\u3000-\u303f]+")

# OMML math namespace — paragraphs that are entirely math expressions
_MATH_PARA_TAG = qn("m:oMathPara")
_MATH_TAG = qn("m:oMath")


def _has_chinese(text: str) -> bool:
    return bool(_CJK_RE.search(text))


def _is_math_element(element) -> bool:
    """Return True if the XML element is an OMML math node."""
    return element.tag in (_MATH_PARA_TAG, _MATH_TAG)


def _paragraph_page_index(para_index: int, doc_paragraph_count: int) -> int:
    """
    Approximate page number for a paragraph using a rough heuristic:
    ~40 paragraphs per page (good enough for --test mode filtering).
    """
    return para_index // 40


def extract_segments(docx_path: Path, test_pages: int | None = None) -> list[dict]:
    """
    Walk all body paragraphs and table cells in *docx_path*.

    Returns a list of segment dicts:
        {
            "id": str,          # unique stable ID
            "source": str,      # Chinese source text
            "location": str,    # human-readable context ("body" | "table")
            "para_index": int,  # index within the paragraph sequence used
            "run_index": int    # run index within the paragraph
        }
    """
    doc = docx.Document(str(docx_path))
    segments: list[dict] = []
    seg_id = 0

    def _process_paragraphs(paragraphs, location_prefix: str, offset: int = 0):
        nonlocal seg_id
        for para_idx, para in enumerate(paragraphs):
            abs_idx = offset + para_idx

            # Apply test-mode page cutoff via heuristic
            if test_pages is not None:
                approx_page = _paragraph_page_index(abs_idx, len(paragraphs))
                if approx_page >= test_pages:
                    break

            # Skip paragraphs whose XML contains only math (oMathPara at top level)
            para_xml = para._element
            if any(_is_math_element(child) for child in para_xml):
                continue

            for run_idx, run in enumerate(para.runs):
                text = run.text
                if not text or not _has_chinese(text):
                    continue
                segments.append(
                    {
                        "id": f"seg_{seg_id:05d}",
                        "source": text,
                        "location": location_prefix,
                        "para_index": abs_idx,
                        "run_index": run_idx,
                    }
                )
                seg_id += 1

    # Body paragraphs
    _process_paragraphs(doc.paragraphs, "body")

    # Table cells
    table_para_offset = len(doc.paragraphs)
    for tbl_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(
                    cell.paragraphs,
                    f"table_{tbl_idx}",
                    offset=table_para_offset,
                )
                table_para_offset += len(cell.paragraphs)

    return segments


def main():
    parser = argparse.ArgumentParser(description="Extract Chinese segments from a DOCX file.")
    parser.add_argument("docx_file", help="Path to the source DOCX file")
    parser.add_argument("--test", action="store_true", help="Enable test mode")
    parser.add_argument(
        "--pages",
        type=int,
        default=None,
        metavar="N",
        help="(test mode) Limit extraction to approximately the first N pages",
    )
    args = parser.parse_args()

    docx_path = Path(args.docx_file)
    if not docx_path.exists():
        print(f"ERROR: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    test_pages = args.pages if args.test else None

    print(f"Extracting Chinese segments from: {docx_path.name}", flush=True)
    if test_pages is not None:
        print(f"  Test mode: ~first {test_pages} page(s)", flush=True)

    segments = extract_segments(docx_path, test_pages=test_pages)

    WORK_DIR.mkdir(parents=True, exist_ok=True)
    with open(EXTRACTED_SEGMENTS, "w", encoding="utf-8") as f:
        json.dump(segments, f, ensure_ascii=False, indent=2)

    print(f"Extracted {len(segments)} segment(s) → {EXTRACTED_SEGMENTS}", flush=True)


if __name__ == "__main__":
    main()
