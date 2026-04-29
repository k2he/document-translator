"""Parse DOCX files and identify paragraphs containing Chinese text."""

import re
from typing import Iterator

from docx import Document
from docx.text.paragraph import Paragraph

# Matches CJK Unified Ideographs and common CJK punctuation/symbols
_CHINESE_RE = re.compile(
    r"[\u4e00-\u9fff"   # CJK Unified Ideographs
    r"\u3400-\u4dbf"   # CJK Extension A
    r"\uf900-\ufaff"   # CJK Compatibility Ideographs
    r"\u3000-\u303f"   # CJK Symbols and Punctuation
    r"\uff00-\uffef]"  # Halfwidth and Fullwidth Forms
)


def has_chinese(text: str) -> bool:
    """Return True if *text* contains at least one Chinese character."""
    return bool(_CHINESE_RE.search(text))


def _paragraph_has_math(paragraph: Paragraph) -> bool:
    """Return True if the paragraph XML contains an OMML math element."""
    # Checking the raw XML string avoids namespace registration complexity.
    return "oMath" in paragraph._element.xml


def _run_text(paragraph: Paragraph) -> str:
    """Return the concatenated text of all <w:r> runs in the paragraph.

    This intentionally excludes OMML elements, which are not runs.
    """
    return "".join(run.text for run in paragraph.runs)


def load_document(path: str) -> Document:
    return Document(path)


def iter_body_paragraphs(doc: Document) -> Iterator[tuple[int, Paragraph]]:
    """Yield (index, paragraph) for every body paragraph with Chinese text."""
    for i, para in enumerate(doc.paragraphs):
        text = _run_text(para)
        if text.strip() and has_chinese(text):
            yield i, para


def iter_table_paragraphs(
    doc: Document,
) -> Iterator[tuple[object, int, Paragraph]]:
    """Yield (cell, para_index, paragraph) for table cells with Chinese text.

    Nested tables are not traversed in v1.
    """
    for table in doc.tables:
        for row in table.rows:
            # Deduplicate merged cells (python-docx exposes them multiple times)
            seen: set[int] = set()
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen:
                    continue
                seen.add(cell_id)
                for j, para in enumerate(cell.paragraphs):
                    text = _run_text(para)
                    if text.strip() and has_chinese(text):
                        yield cell, j, para
