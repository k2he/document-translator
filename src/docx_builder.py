"""Write translated text back into a DOCX document in-place.

Strategy
--------
For each paragraph that has a translation:
  - Place the full translated string in the *first* run's text.
  - Clear the text of all subsequent runs.
  - OMML math elements (<m:oMath>) are siblings of <w:r> nodes in the XML
    and are never touched by this operation, so equations are preserved.
"""

from docx import Document
from docx.text.paragraph import Paragraph


def _apply_to_paragraph(paragraph: Paragraph, translated: str) -> None:
    """Replace paragraph run text with *translated*, preserving first run's formatting."""
    runs = paragraph.runs
    if not runs:
        return
    runs[0].text = translated
    for run in runs[1:]:
        run.text = ""


def apply_body_translations(doc: Document, translations: dict[int, str]) -> None:
    """Apply *translations* (paragraph_index → text) to body paragraphs."""
    for i, para in enumerate(doc.paragraphs):
        if i in translations:
            _apply_to_paragraph(para, translations[i])


def apply_table_translations(
    doc: Document,
    translations: dict[tuple[int, int, int], str],
) -> None:
    """Apply *translations* to table cell paragraphs.

    Key format: (table_index, row_index, para_index) → translated text.
    Merged cells are deduplicated to avoid double-writing.
    """
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            seen: set[int] = set()
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen:
                    continue
                seen.add(cell_id)
                for p_idx, para in enumerate(cell.paragraphs):
                    key = (t_idx, r_idx, p_idx)
                    if key in translations:
                        _apply_to_paragraph(para, translations[key])
